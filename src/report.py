"""
report.py — Generate PDF and DOCX diagnostic reports.
"""

import io
import datetime
import pandas as pd


def generate_pdf_report(chat_history: list, df: pd.DataFrame, anomaly_df: pd.DataFrame = None) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        return b""

    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story  = []

    title_style = ParagraphStyle("Title", parent=styles["Title"], textColor=colors.HexColor("#e94560"), fontSize=24, spaceAfter=6)
    sub_style   = ParagraphStyle("Sub",   parent=styles["Normal"], textColor=colors.HexColor("#555555"), fontSize=10, spaceAfter=20)

    story.append(Paragraph("FailureRadar", title_style))
    story.append(Paragraph(f"Industrial Diagnostic Report — {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}", sub_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#e94560")))
    story.append(Spacer(1, 0.4*cm))

    if not df.empty:
        story.append(Paragraph("Sensor Data Summary", styles["Heading2"]))
        summary_data = [
            ["Metric", "Value"],
            ["Total Machines",  str(df["machine_id"].nunique())],
            ["Total Readings",  str(len(df))],
            ["Critical Events", str(len(df[df["status"] == "critical"]))],
            ["Warning Events",  str(len(df[df["status"] == "warning"]))],
            ["Normal Readings", str(len(df[df["status"] == "normal"]))],
        ]
        t = Table(summary_data, colWidths=[8*cm, 8*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,0), colors.HexColor("#e94560")),
            ("TEXTCOLOR",     (0,0), (-1,0), colors.white),
            ("FONTNAME",      (0,0), (-1,0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [colors.HexColor("#f9f9f9"), colors.white]),
            ("GRID",          (0,0), (-1,-1), 0.5, colors.HexColor("#dddddd")),
            ("PADDING",       (0,0), (-1,-1), 8),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))

    if anomaly_df is not None and "anomaly_flag" in anomaly_df.columns:
        story.append(Paragraph("Isolation Forest Anomaly Detection", styles["Heading2"]))
        flagged = anomaly_df[anomaly_df["anomaly_flag"] == True]
        if not flagged.empty:
            anom_data = [["Machine", "Timestamp", "Temp C", "Vibration", "Score"]]
            for _, row in flagged.head(15).iterrows():
                anom_data.append([str(row.get("machine_id","")), str(row.get("timestamp","")),
                                   str(row.get("temperature_C","")), str(row.get("vibration_mm_s","")),
                                   str(round(row.get("anomaly_score", 0), 3))])
            t2 = Table(anom_data, colWidths=[3.5*cm, 4.5*cm, 2.5*cm, 2.5*cm, 3*cm])
            t2.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
                ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#dddddd")),
                ("FONTSIZE",   (0,0), (-1,-1), 8),
                ("PADDING",    (0,0), (-1,-1), 6),
            ]))
            story.append(t2)
        story.append(Spacer(1, 0.5*cm))

    if chat_history:
        story.append(Paragraph("AI Diagnostic Chat Log", styles["Heading2"]))
        q_style = ParagraphStyle("Q", parent=styles["Normal"], textColor=colors.HexColor("#0f3460"), fontName="Helvetica-Bold", fontSize=10, spaceAfter=4)
        a_style = ParagraphStyle("A", parent=styles["Normal"], textColor=colors.HexColor("#333333"), fontSize=9, spaceAfter=12, leftIndent=10)
        for msg in chat_history:
            if msg["role"] == "user":
                story.append(Paragraph(f"Q: {msg['content']}", q_style))
            else:
                story.append(Paragraph(f"A: {msg['content'].replace(chr(10), '<br/>')}", a_style))

    story.append(Spacer(1, 1*cm))
    footer_style = ParagraphStyle("Footer", parent=styles["Normal"], textColor=colors.HexColor("#999999"), fontSize=8, alignment=TA_CENTER)
    story.append(Paragraph("FailureRadar · RAG + Isolation Forest + Ollama · Built by Puneet Divedi", footer_style))
    doc.build(story)
    return buffer.getvalue()


def generate_docx_report(chat_history: list, df: pd.DataFrame, anomaly_df: pd.DataFrame = None) -> bytes:
    try:
        from docx import Document
        from docx.shared import RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    doc   = Document()
    title = doc.add_heading("FailureRadar — Diagnostic Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    if not df.empty:
        doc.add_heading("Sensor Data Summary", level=1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Metric"
        tbl.rows[0].cells[1].text = "Value"
        for label, val in [
            ("Total Machines",  str(df["machine_id"].nunique())),
            ("Total Readings",  str(len(df))),
            ("Critical Events", str(len(df[df["status"] == "critical"]))),
            ("Warning Events",  str(len(df[df["status"] == "warning"]))),
            ("Normal Readings", str(len(df[df["status"] == "normal"]))),
        ]:
            r = tbl.add_row().cells
            r[0].text = label
            r[1].text = val

    if anomaly_df is not None and "anomaly_flag" in anomaly_df.columns:
        doc.add_heading("Anomaly Detection (Isolation Forest)", level=1)
        flagged = anomaly_df[anomaly_df["anomaly_flag"] == True]
        doc.add_paragraph(f"Total anomalies detected: {len(flagged)}")
        if not flagged.empty:
            atbl = doc.add_table(rows=1, cols=5)
            atbl.style = "Table Grid"
            for i, h in enumerate(["Machine", "Timestamp", "Temp C", "Vibration", "Score"]):
                atbl.rows[0].cells[i].text = h
            for _, row in flagged.head(15).iterrows():
                r = atbl.add_row().cells
                r[0].text = str(row.get("machine_id",""));  r[1].text = str(row.get("timestamp",""))
                r[2].text = str(row.get("temperature_C","")); r[3].text = str(row.get("vibration_mm_s",""))
                r[4].text = str(round(row.get("anomaly_score", 0), 3))

    if chat_history:
        doc.add_heading("AI Diagnostic Chat Log", level=1)
        for msg in chat_history:
            if msg["role"] == "user":
                p = doc.add_paragraph()
                run = p.add_run(f"Engineer: {msg['content']}")
                run.bold = True
                run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
            else:
                p = doc.add_paragraph(f"FailureRadar: {msg['content']}")
                p.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph("FailureRadar · RAG + Isolation Forest + Ollama · Built by Puneet Divedi")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
